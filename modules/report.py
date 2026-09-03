"""
HARMATTAN — Professional multi-format reports (HTML / PDF / DOCX / JSON).
Client-ready consulting-style deliverables.
"""
from __future__ import annotations

import io
from collections import Counter
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Optional

import jinja2

from core.config import VERSION

# ── Jinja2 environment ──────────────────────────────────────────────
_TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"
_JINJA_ENV = jinja2.Environment(
    loader=jinja2.FileSystemLoader(str(_TEMPLATE_DIR)),
    autoescape=jinja2.select_autoescape(["html", "xml"]),
    trim_blocks=True,
    lstrip_blocks=True,
)

# Brand (print-safe consulting palette)
NAVY = "0F172A"
SLATE = "334155"
MUTED = "64748B"
ORANGE = "EA580C"
CYAN = "0D9488"
RED = "DC2626"
GREEN = "16A34A"
BORDER = "E2E8F0"
BG = "F8FAFC"
WHITE = "FFFFFF"


def _meta(network, arp, attack, vuln) -> dict:
    network = network or {}
    arp = arp or {}
    attack = attack or {}
    vuln = vuln or {}
    hosts = arp.get("count") or len(arp.get("hosts") or [])
    exposures = attack.get("total_exposures", 0)
    grade = attack.get("grade", "—")
    score = attack.get("risk_score", 0)
    cves = vuln.get("total_findings", 0)
    return {
        "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "ssid": network.get("ssid") or "—",
        "subnet": network.get("subnet") or "—",
        "gateway": network.get("gateway") or "—",
        "local_ip": network.get("local_ip") or "—",
        "hosts": hosts,
        "exposures": exposures,
        "grade": grade,
        "score": score,
        "cves": cves,
        "roles": arp.get("roles") or {},
        "recommendations": attack.get("recommendations") or [],
    }


def _top_risks(attack: dict, vuln: dict, limit: int = 12) -> list[dict]:
    top_risks: list[dict] = []
    risk_rank = {
        "critique": 0, "critical": 0, "haute": 1, "high": 1,
        "moyenne": 2, "medium": 2, "faible": 3, "low": 3,
    }
    for h in attack.get("hosts", []) or []:
        for e in h.get("exposures") or []:
            top_risks.append({
                "kind": "exposition",
                "ip": h.get("ip"),
                "label": f"Port {e.get('port')} / {e.get('service') or '?'}",
                "risk": e.get("risk") or "moyenne",
                "detail": e.get("recommendation") or e.get("reason") or "",
            })
    for h in vuln.get("hosts", []) or []:
        for s in h.get("services") or []:
            for c in s.get("cves") or []:
                sev = (c.get("severity") or "medium").lower()
                if sev in ("critical", "high", "critique", "haute", "medium", "moyenne"):
                    top_risks.append({
                        "kind": "cve",
                        "ip": h.get("ip"),
                        "label": c.get("id") or "CVE",
                        "risk": sev,
                        "detail": (c.get("description") or "")[:160],
                    })
    top_risks.sort(key=lambda x: risk_rank.get(str(x.get("risk", "")).lower(), 9))
    return top_risks[:limit]


def _exec_text(m: dict) -> str:
    grade = str(m["grade"] or "—")
    if m["exposures"] == 0 and m["cves"] == 0:
        return (
            f"Le segment <b>{escape(str(m['subnet']))}</b> présente un profil globalement maîtrisé "
            f"({m['hosts']} appareil(s), grade <b>{escape(grade)}</b>). "
            "Aucune exposition critique ni CVE corrélée n'a été remarquée dans cette session. "
            "Maintenir le durcissement de base (mises à jour, segmentation, MFA admin)."
        )
    if m["exposures"] >= 5 or str(grade).upper() in ("D", "E", "F"):
        return (
            f"<b>Attention :</b> le réseau <b>{escape(str(m['subnet']))}</b> affiche un niveau de risque élevé "
            f"(grade <b>{escape(grade)}</b>, score {m['score']}/100) avec "
            f"<b>{m['exposures']}</b> exposition(s) et <b>{m['cves']}</b> CVE. "
            "Prioriser la fermeture des services inutiles, le filtrage périmétrique et les correctifs."
        )
    return (
        f"L'audit du segment <b>{escape(str(m['subnet']))}</b> a inventorié "
        f"<b>{m['hosts']}</b> appareil(s). Grade <b>{escape(grade)}</b> ({m['score']}/100) : "
        f"<b>{m['exposures']}</b> exposition(s) de surface d'attaque et "
        f"<b>{m['cves']}</b> corrélation(s) CVE. "
        "Les actions prioritaires figurent dans le top risques et les recommandations ci-dessous."
    )


# ---------------------------------------------------------------------------
# HTML — Jinja2 + Chart.js
# ---------------------------------------------------------------------------
def build_html_report(
    network: Optional[dict],
    arp: Optional[dict],
    nmap: Optional[dict],
    vuln: Optional[dict],
    attack: Optional[dict],
    *,
    title: str = "Rapport d'audit réseau",
    client: str = "",
    operator: str = "NACF / HARMATTAN",
    extras: Optional[list] = None,
) -> str:
    network = network or {}
    arp = arp or {}
    nmap = nmap or {}
    vuln = vuln or {}
    attack = attack or {}
    m = _meta(network, arp, attack, vuln)
    now = m["generated"]
    report_id = f"HMT-{datetime.now().strftime('%Y%m%d-%H%M')}"

    # ── Build data ──

    # Hosts
    hosts = []
    for h in arp.get("hosts", []):
        hosts.append({
            "ip": h.get("ip", ""),
            "mac": h.get("mac") or "—",
            "vendor": h.get("vendor") or "—",
            "hostname": h.get("hostname") or "—",
            "role": h.get("role") or "—",
            "os_hint": h.get("os_hint") or "—",
            "ports_str": ", ".join(str(p) for p in h.get("open_ports", [])) or "—",
        })

    # Nmap sections
    nmap_sections = []
    for h in nmap.get("hosts", []):
        ports = []
        for p in h.get("ports", []):
            if p.get("state") != "open":
                continue
            ports.append({
                "port": str(p.get("port")),
                "proto": p.get("protocol", ""),
                "service": p.get("service", ""),
                "product": (p.get("product") or "") + " " + (p.get("version") or ""),
            })
        nmap_sections.append({
            "ip": h.get("ip", ""),
            "os": h["os_matches"][0]["name"] if h.get("os_matches") else "—",
            "ports": ports,
        })

    # Attack surface rows
    attack_rows = []
    for h in attack.get("hosts", []):
        for e in h.get("exposures", []):
            attack_rows.append({
                "ip": h.get("ip", ""),
                "role": h.get("role", ""),
                "port": e.get("port"),
                "service": e.get("service", ""),
                "risk": e.get("risk", ""),
                "reco": e.get("recommendation", "—"),
            })

    # CVE rows
    cve_rows = []
    for h in vuln.get("hosts", []):
        for s in h.get("services", []):
            for c in s.get("cves", []):
                cve_rows.append({
                    "ip": h.get("ip", ""),
                    "product": f"{s.get('product','')} {s.get('version','')}".strip(),
                    "id": c.get("id", ""),
                    "url": c.get("url", ""),
                    "score": c.get("score") if c.get("score") is not None else "—",
                    "severity": c.get("severity", ""),
                    "description": (c.get("description") or "")[:220],
                })

    # Recommendations
    recommendations = m["recommendations"][:20]

    # Risk labels for chart
    risk_counts: Counter[str] = Counter()
    for r in _top_risks(attack, vuln, 50):
        risk = r.get("risk", "info").lower()
        if risk in ("critical", "critique"):
            risk_counts["Critique"] += 1
        elif risk in ("high", "haute", "élevée", "elevee"):
            risk_counts["Haute"] += 1
        elif risk in ("medium", "moyenne"):
            risk_counts["Moyenne"] += 1
        else:
            risk_counts["Faible"] += 1
    risk_labels = [{"label": k, "count": v} for k, v in risk_counts.most_common()]

    # Host chart data (ports per host)
    host_chart = []
    for h in arp.get("hosts", [])[:10]:
        n_ports = len(h.get("open_ports", []))
        if n_ports > 0:
            host_chart.append({"ip": h.get("ip", "?"), "ports": n_ports})

    # Risk bar data
    risk_bar = {
        "crit": risk_counts.get("Critique", 0),
        "high": risk_counts.get("Haute", 0),
        "med": risk_counts.get("Moyenne", 0),
        "low": risk_counts.get("Faible", 0),
    }

    # Grade class
    g = str(m["grade"] or "").upper()
    if g in ("D", "E", "F"):
        grade_class = "bad"
    elif g in ("C",):
        grade_class = "mid"
    else:
        grade_class = "ok"

    top_risks_data = _top_risks(attack, vuln)

    try:
        template = _JINJA_ENV.get_template("report_audit.html")
        return template.render(
            title=title,
            client=client or "Interne",
            operator=operator,
            generated=now,
            report_id=report_id,
            version=VERSION,
            ssid=m["ssid"],
            subnet=m["subnet"],
            gateway=m["gateway"],
            gateway_short=str(m["gateway"])[:14] if m["gateway"] else "—",
            hosts_count=m["hosts"],
            exposures_count=m["exposures"],
            cves_count=m["cves"],
            grade=str(m["grade"]),
            score=m["score"],
            grade_class=grade_class,
            exec_text=_exec_text(m),
            top_risks=top_risks_data,
            risk_labels=risk_labels,
            risk_bar=risk_bar,
            host_chart=host_chart,
            roles=m["roles"],
            recommendations=recommendations,
            hosts=hosts,
            attack_hosts_total=attack.get("total_hosts", 0),
            attack_rows=attack_rows,
            nmap_sections=nmap_sections,
            cve_rows=cve_rows,
            extras=extras or [],
        )
    except Exception as exc:
        # Fallback: basic HTML
        parts = [f"<h1>{escape(title)}</h1>",
                 f"<p>Généré {escape(now)} · {report_id}</p>",
                 f"<p>Hôtes: {m['hosts']} · Expositions: {m['exposures']} · CVE: {m['cves']}</p>",
                 f"<p><em>Template Jinja2: {escape(str(exc))[:200]}</em></p>"]
        return "<!DOCTYPE html><html><body>" + "".join(parts) + "</body></html>"


def build_json_report(
    network: Optional[dict],
    arp: Optional[dict],
    nmap: Optional[dict],
    vuln: Optional[dict],
    attack: Optional[dict],
    **kwargs,
) -> dict:
    m = _meta(network, arp, attack, vuln)
    return {
        "meta": {
            "generated": datetime.now().isoformat(timespec="seconds"),
            "tool": "HARMATTAN",
            "version": VERSION,
            "report_type": "network_audit",
            "classification": "CONFIDENTIEL",
            **{k: v for k, v in kwargs.items() if v},
        },
        "summary": m,
        "top_risks": _top_risks(attack or {}, vuln or {}),
        "network": network,
        "arp": arp,
        "nmap": nmap,
        "vuln": vuln,
        "attack": attack,
    }


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def build_pdf_report(
    network: Optional[dict],
    arp: Optional[dict],
    nmap: Optional[dict],
    vuln: Optional[dict],
    attack: Optional[dict],
    *,
    title: str = "Rapport d'audit réseau",
    client: str = "",
    operator: str = "NACF / HARMATTAN",
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
        HRFlowable,
        KeepTogether,
    )

    network = network or {}
    arp = arp or {}
    nmap = nmap or {}
    vuln = vuln or {}
    attack = attack or {}
    m = _meta(network, arp, attack, vuln)
    report_id = f"HMT-{datetime.now().strftime('%Y%m%d-%H%M')}"
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=18 * mm,
        bottomMargin=16 * mm,
        title=f"HARMATTAN — {title}",
        author=operator,
        subject="Rapport d'audit réseau confidentiel",
    )

    navy = colors.HexColor("#0F172A")
    orange = colors.HexColor("#EA580C")
    muted = colors.HexColor("#64748B")
    border = colors.HexColor("#E2E8F0")
    bg = colors.HexColor("#F8FAFC")
    soft = colors.HexColor("#F1F5F9")

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CoverTitle", fontSize=20, leading=24, textColor=navy, spaceAfter=6, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="H2", fontSize=11, leading=14, textColor=navy, spaceBefore=14, spaceAfter=8, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="Body", fontSize=9, leading=12, textColor=colors.HexColor("#334155"), alignment=TA_JUSTIFY))
    styles.add(ParagraphStyle(name="Muted", fontSize=8, leading=10, textColor=muted))
    styles.add(ParagraphStyle(name="Brand", fontSize=8, textColor=navy, fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="Cell", fontSize=7.5, leading=9.5, textColor=navy))
    styles.add(ParagraphStyle(name="CellHead", fontSize=7.5, leading=9, textColor=colors.HexColor("#475569"), fontName="Helvetica-Bold"))
    styles.add(ParagraphStyle(name="Classif", fontSize=8, textColor=orange, fontName="Helvetica-Bold", alignment=TA_RIGHT))

    story = []

    # Header brand strip
    head = Table(
        [[
            Paragraph("HARMATTAN · NETWORK INTELLIGENCE", styles["Brand"]),
            Paragraph("CONFIDENTIEL", styles["Classif"]),
        ]],
        colWidths=[120 * mm, 50 * mm],
    )
    head.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), soft),
        ("BOX", (0, 0), (-1, -1), 0.5, border),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(head)
    story.append(Spacer(1, 10))
    story.append(Paragraph(escape(title), styles["CoverTitle"]))
    story.append(Paragraph(
        f"<b>Client:</b> {escape(client or 'Interne')} &nbsp;·&nbsp; "
        f"<b>Opérateur:</b> {escape(operator)} &nbsp;·&nbsp; "
        f"<b>Date:</b> {escape(m['generated'])} &nbsp;·&nbsp; "
        f"<b>Réf:</b> {report_id} &nbsp;·&nbsp; v{VERSION}",
        styles["Muted"],
    ))
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=2, color=navy, spaceAfter=8))
    story.append(Paragraph(
        "Document réservé aux parties autorisées. Audit réalisé dans un cadre légal explicite.",
        styles["Muted"],
    ))
    story.append(Spacer(1, 8))

    # Network context table
    ctx = [
        [Paragraph("<b>SSID</b>", styles["CellHead"]), Paragraph(escape(str(m["ssid"])), styles["Cell"]),
         Paragraph("<b>Subnet</b>", styles["CellHead"]), Paragraph(escape(str(m["subnet"])), styles["Cell"])],
        [Paragraph("<b>Gateway</b>", styles["CellHead"]), Paragraph(escape(str(m["gateway"])), styles["Cell"]),
         Paragraph("<b>IP locale</b>", styles["CellHead"]), Paragraph(escape(str(m["local_ip"])), styles["Cell"])],
    ]
    t = Table(ctx, colWidths=[25 * mm, 55 * mm, 25 * mm, 55 * mm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), bg),
        ("BOX", (0, 0), (-1, -1), 0.5, border),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, border),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)

    # 1. Executive Summary
    story.append(Paragraph("1. Synthèse exécutive", styles["H2"]))
    kpi = [[
        Paragraph(
            f"<font size='26' color='#0F172A'><b>{escape(str(m['grade']))}</b></font>"
            f"<br/><font size='8' color='#64748B'>Grade · {m['score']}/100</font>",
            styles["Cell"],
        ),
        Paragraph(f"<font size='14'><b>{m['hosts']}</b></font><br/><font size='8' color='#64748B'>Appareils</font>", styles["Cell"]),
        Paragraph(f"<font size='14'><b>{m['exposures']}</b></font><br/><font size='8' color='#64748B'>Expositions</font>", styles["Cell"]),
        Paragraph(f"<font size='14'><b>{m['cves']}</b></font><br/><font size='8' color='#64748B'>CVE</font>", styles["Cell"]),
    ]]
    kt = Table(kpi, colWidths=[42 * mm, 40 * mm, 40 * mm, 40 * mm])
    kt.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF7ED")),
        ("BOX", (0, 0), (-1, -1), 1, orange),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
    ]))
    story.append(kt)
    story.append(Spacer(1, 8))
    story.append(Paragraph(
        f"L'audit a identifié <b>{m['hosts']}</b> appareil(s) sur <b>{escape(str(m['subnet']))}</b>, "
        f"avec <b>{m['exposures']}</b> exposition(s) et <b>{m['cves']}</b> corrélation(s) CVE. "
        f"Grade global de risque : <b>{escape(str(m['grade']))}</b> ({m['score']}/100).",
        styles["Body"],
    ))
    # Executive narrative
    story.append(Paragraph(_exec_text(m), styles["Body"]))

    # 2. Top risks
    story.append(Paragraph("2. Top risques prioritaires", styles["H2"]))
    trh = [Paragraph(x, styles["CellHead"]) for x in ["#", "IP", "Type", "Élément", "Sev", "Détail"]]
    tr_rows = [trh]
    for i, r in enumerate(_top_risks(attack, vuln, 10), 1):
        tr_rows.append([
            Paragraph(str(i), styles["Cell"]),
            Paragraph(escape(str(r.get("ip") or "—")), styles["Cell"]),
            Paragraph(escape(str(r.get("kind") or "")), styles["Cell"]),
            Paragraph(escape(str(r.get("label") or "")[:28]), styles["Cell"]),
            Paragraph(escape(str(r.get("risk") or "")), styles["Cell"]),
            Paragraph(escape(str(r.get("detail") or "—")[:55]), styles["Cell"]),
        ])
    if len(tr_rows) == 1:
        tr_rows.append([Paragraph("Aucun risque prioritaire", styles["Cell"])] + [Paragraph("", styles["Cell"])] * 5)
    trt = Table(tr_rows, colWidths=[10 * mm, 28 * mm, 22 * mm, 40 * mm, 20 * mm, 52 * mm])
    trt.setStyle(_table_style())
    story.append(trt)

    # 3. Recommendations
    if m["recommendations"]:
        story.append(Paragraph("3. Recommandations prioritaires", styles["H2"]))
        for r in m["recommendations"][:15]:
            story.append(Paragraph(f"• {escape(r)}", styles["Body"]))

    # 4. Host inventory
    story.append(Paragraph("4. Inventaire des hôtes", styles["H2"]))
    header = [Paragraph(x, styles["CellHead"]) for x in ["IP", "MAC", "Vendor", "Hostname", "Rôle", "Ports"]]
    rows = [header]
    for h in (arp.get("hosts") or [])[:80]:
        ports = ", ".join(str(p) for p in h.get("open_ports", [])[:8]) or "—"
        rows.append([
            Paragraph(escape(h.get("ip") or ""), styles["Cell"]),
            Paragraph(escape(h.get("mac") or "—"), styles["Cell"]),
            Paragraph(escape((h.get("vendor") or "—")[:24]), styles["Cell"]),
            Paragraph(escape((h.get("hostname") or "—")[:20]), styles["Cell"]),
            Paragraph(escape(h.get("role") or "—"), styles["Cell"]),
            Paragraph(escape(ports[:40]), styles["Cell"]),
        ])
    if len(rows) == 1:
        rows.append([Paragraph("Aucun hôte", styles["Cell"])] + [Paragraph("", styles["Cell"])] * 5)
    ht = Table(rows, colWidths=[28 * mm, 32 * mm, 32 * mm, 28 * mm, 22 * mm, 30 * mm])
    ht.setStyle(_table_style())
    story.append(ht)

    # 5. Attack surface
    story.append(Paragraph("5. Surface d'attaque", styles["H2"]))
    ah = [Paragraph(x, styles["CellHead"]) for x in ["IP", "Port", "Service", "Risque", "Reco"]]
    arows = [ah]
    for h in attack.get("hosts") or []:
        for e in h.get("exposures") or []:
            arows.append([
                Paragraph(escape(h.get("ip") or ""), styles["Cell"]),
                Paragraph(str(e.get("port") or ""), styles["Cell"]),
                Paragraph(escape((e.get("service") or "")[:18]), styles["Cell"]),
                Paragraph(escape(e.get("risk") or ""), styles["Cell"]),
                Paragraph(escape((e.get("recommendation") or "—")[:50]), styles["Cell"]),
            ])
            if len(arows) > 60:
                break
        if len(arows) > 60:
            break
    if len(arows) == 1:
        arows.append([Paragraph("Aucune exposition", styles["Cell"])] + [Paragraph("", styles["Cell"])] * 4)
    at = Table(arows, colWidths=[28 * mm, 16 * mm, 28 * mm, 22 * mm, 72 * mm])
    at.setStyle(_table_style())
    story.append(at)

    # 6. Nmap
    story.append(Paragraph("6. Scans Nmap (extrait)", styles["H2"]))
    nh = [Paragraph(x, styles["CellHead"]) for x in ["IP", "Port", "Service", "Produit / Version"]]
    nrows = [nh]
    for h in nmap.get("hosts") or []:
        for p in h.get("ports") or []:
            if p.get("state") != "open":
                continue
            nrows.append([
                Paragraph(escape(h.get("ip") or ""), styles["Cell"]),
                Paragraph(f"{p.get('port')}/{p.get('protocol')}", styles["Cell"]),
                Paragraph(escape(p.get("service") or ""), styles["Cell"]),
                Paragraph(escape(f"{p.get('product') or ''} {p.get('version') or ''}".strip()[:40]), styles["Cell"]),
            ])
            if len(nrows) > 50:
                break
        if len(nrows) > 50:
            break
    if len(nrows) == 1:
        nrows.append([Paragraph("Aucun résultat nmap", styles["Cell"])] + [Paragraph("", styles["Cell"])] * 3)
    nt = Table(nrows, colWidths=[30 * mm, 22 * mm, 30 * mm, 84 * mm])
    nt.setStyle(_table_style())
    story.append(nt)

    # 7. CVE
    story.append(Paragraph("7. Vulnérabilités CVE (extrait)", styles["H2"]))
    vh = [Paragraph(x, styles["CellHead"]) for x in ["IP", "Produit", "CVE", "Score", "Sev"]]
    vrows = [vh]
    for h in vuln.get("hosts") or []:
        for s in h.get("services") or []:
            for c in s.get("cves") or []:
                vrows.append([
                    Paragraph(escape(h.get("ip") or ""), styles["Cell"]),
                    Paragraph(escape(f"{s.get('product','')} {s.get('version','')}"[:28]), styles["Cell"]),
                    Paragraph(escape(c.get("id") or ""), styles["Cell"]),
                    Paragraph(str(c.get("score") if c.get("score") is not None else "—"), styles["Cell"]),
                    Paragraph(escape(c.get("severity") or ""), styles["Cell"]),
                ])
                if len(vrows) > 40:
                    break
            if len(vrows) > 40:
                break
        if len(vrows) > 40:
            break
    if len(vrows) == 1:
        vrows.append([Paragraph("Aucune CVE", styles["Cell"])] + [Paragraph("", styles["Cell"])] * 4)
    vt = Table(vrows, colWidths=[28 * mm, 50 * mm, 35 * mm, 18 * mm, 25 * mm])
    vt.setStyle(_table_style())
    story.append(vt)

    # 8. Methodology
    story.append(Paragraph("8. Méthodologie & limitations", styles["H2"]))
    story.append(Paragraph(
        "Découverte ARP + fingerprinting, scans nmap, scoring de surface d'attaque, "
        "corrélation NVD. Les hôtes silencieux peuvent être absents. Les CVE dépendent "
        "de la précision product/version. Aucune exploitation destructive n'est incluse.",
        styles["Body"],
    ))
    story.append(Spacer(1, 14))
    story.append(HRFlowable(width="100%", thickness=1, color=navy, spaceAfter=6))
    story.append(Paragraph(
        f"HARMATTAN v{VERSION} · NACF · {escape(m['generated'])} · {report_id} · Confidentiel",
        styles["Muted"],
    ))

    def _footer(canvas, doc_):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#0F172A"))
        canvas.setLineWidth(0.6)
        w, h = A4
        canvas.line(15 * mm, h - 12 * mm, w - 15 * mm, h - 12 * mm)
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(muted)
        canvas.drawString(15 * mm, h - 10 * mm, "HARMATTAN — Rapport d'audit réseau")
        canvas.drawRightString(w - 15 * mm, h - 10 * mm, "CONFIDENTIEL")
        canvas.line(15 * mm, 12 * mm, w - 15 * mm, 12 * mm)
        canvas.drawString(15 * mm, 8 * mm, report_id)
        canvas.drawRightString(w - 15 * mm, 8 * mm, f"Page {doc_.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


def _table_style():
    from reportlab.lib import colors
    from reportlab.platypus import TableStyle
    return TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F1F5F9")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#475569")),
        ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#E2E8F0")),
        ("INNERGRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#E2E8F0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FAFBFC")]),
    ])


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def build_docx_report(
    network: Optional[dict],
    arp: Optional[dict],
    nmap: Optional[dict],
    vuln: Optional[dict],
    attack: Optional[dict],
    *,
    title: str = "Rapport d'audit réseau",
    client: str = "",
    operator: str = "NACF / HARMATTAN",
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    network = network or {}
    arp = arp or {}
    nmap = nmap or {}
    vuln = vuln or {}
    attack = attack or {}
    m = _meta(network, arp, attack, vuln)
    report_id = f"HMT-{datetime.now().strftime('%Y%m%d-%H%M')}"

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    brand = doc.add_paragraph()
    run = brand.add_run("HARMATTAN · NETWORK INTELLIGENCE")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    run2 = brand.add_run("    CONFIDENTIEL")
    run2.bold = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)

    h = doc.add_heading(title, level=0)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    meta_run = p.add_run(
        f"Client: {client or 'Interne'}  ·  Opérateur: {operator}  ·  "
        f"Date: {m['generated']}  ·  Réf: {report_id}  ·  v{VERSION}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    conf = doc.add_paragraph()
    r = conf.add_run("Document confidentiel — usage réservé au destinataire autorisé.")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)

    # 1. Executive
    doc.add_heading("1. Synthèse exécutive", level=1)
    doc.add_paragraph(
        f"Grade: {m['grade']} ({m['score']}/100) — "
        f"{m['hosts']} appareil(s), {m['exposures']} exposition(s), {m['cves']} CVE. "
        f"Segment {m['subnet']} · Gateway {m['gateway']} · SSID {m['ssid']}."
    )
    # Executive narrative from _exec_text (strip HTML tags for DOCX)
    exec_str = _exec_text(m).replace("<b>", "").replace("</b>", "")
    doc.add_paragraph(exec_str)

    # 2. Top risks
    doc.add_heading("2. Top risques prioritaires", level=1)
    tr = doc.add_table(rows=1, cols=5)
    tr.style = "Table Grid"
    for i, name in enumerate(["IP", "Type", "Élément", "Sévérité", "Détail"]):
        tr.rows[0].cells[i].text = name
        _shade_cell(tr.rows[0].cells[i], "F1F5F9")
    for risk in _top_risks(attack, vuln, 12):
        row = tr.add_row().cells
        row[0].text = str(risk.get("ip") or "")
        row[1].text = str(risk.get("kind") or "")
        row[2].text = str(risk.get("label") or "")[:40]
        row[3].text = str(risk.get("risk") or "")
        row[4].text = str(risk.get("detail") or "")[:80]

    # 3. Recommendations
    if m["recommendations"]:
        doc.add_heading("3. Recommandations prioritaires", level=1)
        for rec in m["recommendations"][:15]:
            doc.add_paragraph(rec, style="List Bullet")

    # 4. Hosts
    doc.add_heading("4. Inventaire des hôtes", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, name in enumerate(["IP", "MAC", "Vendor", "Hostname", "Rôle", "Ports"]):
        hdr[i].text = name
        _shade_cell(hdr[i], "F1F5F9")
    for hst in (arp.get("hosts") or [])[:100]:
        row = table.add_row().cells
        row[0].text = hst.get("ip") or ""
        row[1].text = hst.get("mac") or ""
        row[2].text = (hst.get("vendor") or "")[:40]
        row[3].text = hst.get("hostname") or ""
        row[4].text = hst.get("role") or ""
        row[5].text = ", ".join(str(p) for p in hst.get("open_ports", [])[:10])

    # 5. Attack surface
    doc.add_heading("5. Surface d'attaque", level=1)
    at = doc.add_table(rows=1, cols=5)
    at.style = "Table Grid"
    for i, name in enumerate(["IP", "Port", "Service", "Risque", "Remédiation"]):
        at.rows[0].cells[i].text = name
        _shade_cell(at.rows[0].cells[i], "F1F5F9")
    n = 0
    for hst in attack.get("hosts") or []:
        for e in hst.get("exposures") or []:
            row = at.add_row().cells
            row[0].text = hst.get("ip") or ""
            row[1].text = str(e.get("port") or "")
            row[2].text = e.get("service") or ""
            row[3].text = e.get("risk") or ""
            row[4].text = (e.get("recommendation") or "")[:120]
            n += 1
            if n >= 80:
                break
        if n >= 80:
            break

    # 6. CVE
    doc.add_heading("6. CVE", level=1)
    vt = doc.add_table(rows=1, cols=5)
    vt.style = "Table Grid"
    for i, name in enumerate(["IP", "Produit", "CVE", "Score", "Sévérité"]):
        vt.rows[0].cells[i].text = name
        _shade_cell(vt.rows[0].cells[i], "F1F5F9")
    n = 0
    for hst in vuln.get("hosts") or []:
        for s in hst.get("services") or []:
            for c in s.get("cves") or []:
                row = vt.add_row().cells
                row[0].text = hst.get("ip") or ""
                row[1].text = f"{s.get('product','')} {s.get('version','')}"[:40]
                row[2].text = c.get("id") or ""
                row[3].text = str(c.get("score") if c.get("score") is not None else "—")
                row[4].text = c.get("severity") or ""
                n += 1
                if n >= 50:
                    break
            if n >= 50:
                break
        if n >= 50:
            break

    # 7. Methodology
    doc.add_heading("7. Méthodologie", level=1)
    doc.add_paragraph(
        "ARP discovery + fingerprinting, nmap, attack surface scoring, corrélation NVD. "
        "Livrable non destructif. Usage strictement autorisé. "
        "Les résultats reflètent l'état du réseau au moment du scan uniquement."
    )
    foot = doc.add_paragraph()
    fr = foot.add_run(f"HARMATTAN v{VERSION} · NACF · {m['generated']} · {report_id} · Confidentiel")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)
